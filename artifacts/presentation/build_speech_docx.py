# -*- coding: utf-8 -*-
"""
Илтгэлийн ярианы текстийг Word файл (.docx) болгож үүсгэнэ.
Эх сурвалж: build_workflow_pptx.py-аас SPEAKER_NOTES dict-ийг импорт хийнэ.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_workflow_pptx import SPEAKER_NOTES

SLIDE_TITLES = {
    1:  "Slide 1 — Гарчиг",
    2:  "Slide 2 — Үндэслэл",
    3:  "Slide 3 — Харьцуулсан судалгаа",
    4:  "Slide 4 — Зорилго ба архитектур",
    5:  "Slide 5 — Зочны хэсэг (3 mobile screenshot)",
    6:  "Slide 6 — Ажилтны самбар (2 desktop screenshot)",
    7:  "Slide 6b — Нэмэлт модулиуд (3 desktop screenshot)",
    8:  "Slide 7 — Захиалгын 7 алхамт урсгал",
    9:  "Slide 8 — Multi-guest QR session",
    10: "Slide 9 — Realtime захиалгын дамжуулалт",
    11: "Slide 10 — Менежерийн 9 модуль",
    12: "Slide 11 — AR давуу талууд",
    13: "Slide 11b — AR жинхэнэ ажиллагаа (2 screenshot)",
    14: "Slide 12 — AR ажиллах 6 алхамт явц",
    15: "Slide 13 — Туршилт ба үр дүн",
    16: "Slide 14 — Дүгнэлт + Q&A",
}

PRIMARY = RGBColor(0x00, 0x46, 0x7F)
ACCENT  = RGBColor(0x46, 0x82, 0xB4)
GOLD    = RGBColor(0xB8, 0x8A, 0x00)
GRAY    = RGBColor(0x60, 0x60, 0x60)


def add_heading(doc, text, level=1, color=PRIMARY):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = "DejaVu Sans"
    if level == 0:
        run.font.size = Pt(20)
    elif level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return h


def add_para(doc, text, *, size=11, bold=False, italic=False, color=None,
             align=None, font="DejaVu Sans"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return p


def build_docx():
    doc = Document()

    # Хуудасны margin
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── Гарчиг ───
    add_heading(doc, "ХАМГААЛАЛТЫН ИЛТГЭЛИЙН ТЕКСТ", level=0, color=PRIMARY)
    add_para(doc,
             "QR код дээр суурилсан ресторан захиалгын систем — Хан Гарид",
             size=14, bold=True, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
             "Эрдэнэтөгсийн Очбадрах   ·   2026 он   ·   Өмнөговь ТДС",
             size=11, italic=True, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ─── Тэмдэглэгээ ───
    add_heading(doc, "Тэмдэглэгээний тайлбар", level=2)
    add_para(doc, "[ДЭЛГЭРЭНГҮЙ ЯРИХ] — энэ слайдын дээр удаан, дэлгэрэнгүй ярина (60-100 сек).",
             size=10, italic=True)
    add_para(doc, "[ДУРДАЖ ӨНГӨРӨХ] — товч, гол санааг л дамжуулна (30-60 сек).",
             size=10, italic=True)
    doc.add_paragraph()

    # ─── Слайд бүрд харгалзах текст ───
    for idx in sorted(SPEAKER_NOTES.keys()):
        title = SLIDE_TITLES.get(idx, f"Slide {idx}")
        add_heading(doc, title, level=1)
        text = SPEAKER_NOTES[idx].strip()
        # Эхний мөр нь [ДЭЛГЭРЭНГҮЙ ...] эсвэл [ДУРДАЖ ...] байх магадлал өндөр
        lines = text.split("\n")
        first_line = lines[0].strip() if lines else ""
        body = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
        if first_line.startswith("["):
            add_para(doc, first_line, size=11, bold=True, color=GOLD)
        else:
            body = text
        if body:
            # Хоёр давхар хоосон мөрийг paragraph болгож хуваана
            for para in body.split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                add_para(doc, para, size=11)
        doc.add_paragraph()

    # ─── Q&A хүснэгт ───
    doc.add_page_break()
    add_heading(doc, "Бэлдсэн Q&A — магадлалтай асуултууд",
                level=1, color=PRIMARY)
    qa = [
        ("Олон хэрэглэгч ачаалал нэмэгдвэл хэрхэх вэ?",
         "Horizontal scale + Redis adapter ашиглан Socket.IO олон node "
         "хооронд sync хийнэ. Load balancer-ийн ард Express instance "
         "олныг ажиллуулна."),
        ("Системийн аюулгүй байдал хэрхэн хангагдсан вэ?",
         "JWT баталгаажуулалт, bcrypt password hashing (cost=10), "
         "role-based middleware (manager/chef/cashier), CORS whitelist, "
         "rate limiting, SQL injection-ээс хамгаалах Drizzle ORM "
         "parameterized query."),
        ("Систем дотор ямар нэгэн bug, дутагдал бий юу?",
         "8 functional test бүгд PASS, гэвч энэ нь демо орчны "
         "preliminary хэмжээс. Бодит ресторанд явуулах field study "
         "ирээдүйн ажил. Edge case-уудыг бүрэн нийгмийн орчинд шалгах "
         "шаардлагатай."),
        ("Системийг хэрэгжүүлэхэд хичнээн зардал шаардагдах вэ?",
         "Open-source бүх компонент. Демо хувилбар одоо локал Windows "
         "компьютер дээр ажиллаж, Cloudflare quick tunnel (үнэгүй) "
         "ашиглан гадны хандалттай. Бодит ресторанд нэвтрүүлэхэд "
         "ердийн desktop эсвэл mini-PC (~1 сая ₮) хангалттай, "
         "интернэт + Cloudflare Tunnel үнэгүй. Зочид өөрсдийн утсаа "
         "ашиглах тул нэмэлт тоног төхөөрөмж шаардлагагүй."),
        ("AR функц дэмжихгүй утсанд яах вэ?",
         "WebGL 3D rotate fallback автоматаар идэвхждэг. Зочин "
         "хоолыг 360° эргүүлж харах боломжтой. ARCore/ARKit "
         "шаардахгүй учир ихэнх ухаалаг гар утсанд ажиллана."),
        ("Multi-guest session-д зочид хоорондоо хэрхэн зөрчилдөхгүй вэ?",
         "Хамтын сагсанд item нэмэхэд session_participants хүснэгтэд "
         "added_by_participant_id хадгалагдана. Зөвхөн item нэмсэн "
         "зочин өөрийнхөө мөрийг устгах боломжтой, бусдынхыг "
         "устгахгүй. Захиалга илгээх эрх manager session creator-д бий."),
        ("QPay-тай хэрхэн холбогддог вэ?",
         "Хамгаалалтын демод QPay sandbox API ашиглан төлбөрийн "
         "QR кодыг үүсгэж, callback URL-аар баталгаажуулах flow "
         "хийсэн. Production-д QPay merchant эрхтэй холбоход "
         "тохиргооны 2-3 хувьсагч өөрчлөхөд хангалттай."),
        ("Ирээдүйд ямар сайжруулалт хийх төлөвлөгөөтэй вэ?",
         "1. Loyalty point систем, 2. Reservation/booking функц, "
         "3. Олон ресторанд multi-tenant дэмжлэг, 4. Зочдын reorder "
         "санал болгол (ML-based recommendation), 5. Voice ordering "
         "Mongolian NLP, 6. POS интеграц (нягтлан бодох)."),
    ]
    for i, (q, a) in enumerate(qa, start=1):
        add_para(doc, f"А{i}. {q}", size=11, bold=True, color=ACCENT)
        add_para(doc, a, size=11)
        doc.add_paragraph()

    # ─── Цагийн хуваарь ───
    doc.add_page_break()
    add_heading(doc, "Цагийн хуваарь (нийт 14-15 мин)",
                level=1, color=PRIMARY)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Хэсэг"
    hdr[1].text = "Слайдууд"
    hdr[2].text = "Ойролцоо цаг"
    schedule = [
        ("Үндэс ба зорилго",         "1-4",     "~4.5 мин"),
        ("Дэлгэцийн танилцуулга",    "5-6b",    "~2.5 мин"),
        ("WORKFLOW + Multi-guest + Realtime (гол хэсэг)", "7-9",   "~5 мин"),
        ("Менежер ба AR",            "10-12",   "~2 мин"),
        ("Туршилт ба хаалт",         "13-15",   "~2 мин"),
        ("Q&A (нөөц)",               "—",       "3-5 мин"),
    ]
    for row in schedule:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = val

    doc.add_paragraph()
    add_para(doc,
             "Зөвлөмж: Slide 7 (7 алхамт workflow), Slide 8 (multi-guest), "
             "Slide 9 (realtime) — эдгээр 3 слайд бол хамгийн чухал. "
             "Бусад слайдыг шаардлагатай бол хурдан өнгөрөөж болно.",
             size=10, italic=True, color=GRAY)

    # ─── Сэв сар ───
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "speech.docx")
    doc.save(out)
    try:
        print(f"[OK] Word файл амжилттай үүсгэгдлээ: {out}")
        print(f"     Слайдын тоо: {len(SPEAKER_NOTES)}")
    except UnicodeEncodeError:
        pass
    return out


if __name__ == "__main__":
    build_docx()
